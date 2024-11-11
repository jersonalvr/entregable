import streamlit as st
import time
import pandas as pd
import calendar
from datetime import datetime
import os
import json
import matplotlib.pyplot as plt
from openpyxl.styles import Font
from collections import defaultdict
from docx import Document
from docx.shared import Pt, Inches, Mm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import pdfplumber
import re
import requests
from io import BytesIO
import tempfile
import zipfile
import pyperclip
from st_copy_to_clipboard import st_copy_to_clipboard
import shutil
from contextlib import contextmanager
from docxtpl import DocxTemplate, InlineImage
import logging
import platform
import subprocess

# Importaciones específicas de Windows
if platform.system() == 'Windows':
    import win32com.client as win32
    import pythoncom
else:
    # No importar pythoncom en sistemas no-Windows
    pass

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Configuración de la página
st.set_page_config(
    page_title="Genera tu Entregable",
    page_icon="🎣",
    layout="wide"
)

# Diccionario para meses en español
MESES_ES = {
    1: "Enero",
    2: "Febrero",
    3: "Marzo",
    4: "Abril",
    5: "Mayo",
    6: "Junio",
    7: "Julio",
    8: "Agosto",
    9: "Septiembre",
    10: "Octubre",
    11: "Noviembre",
    12: "Diciembre",
}

# Determinar la ruta base de la aplicación
base_dir = os.path.dirname(os.path.abspath(__file__))

# Context manager para manejar la inicialización y desinicialización de COM
@contextmanager
def com_handler():
    pythoncom.CoInitialize()
    try:
        yield
    finally:
        pythoncom.CoUninitialize()

# Función para cargar especies desde un archivo JSON
def cargar_especies(archivo="especies.json"):
    plantilla_path = os.path.join(base_dir, archivo)
    if not os.path.exists(plantilla_path):
        st.error(f"El archivo {archivo} no se encontró.")
        return {}
    try:
        with open(plantilla_path, "r", encoding="utf-8") as f:
            especies = json.load(f)
        return especies
    except Exception as e:
        st.error(f"Error al leer el archivo {archivo}: {e}")
        return {}

# Cargar las ciudades desde el archivo JSON
def cargar_ciudades():
    with open(os.path.join(base_dir, "ciudades.json"), "r", encoding="utf-8") as f:
        return json.load(f)

# Función para obtener datos desde la API de SUNAT
def obtener_datos_sunat(dni):
    apisnet_key = st.secrets["APISNET"]["key"]
    url = f"https://api.apis.net.pe/v2/sunat/dni?numero={dni}&token={apisnet_key}"
    try:
        response = requests.get(url)
        if response.status_code == 200:
            data = response.json()
            nombres = f"{data.get('nombres', '')} {data.get('apellidoPaterno', '')} {data.get('apellidoMaterno', '')}".strip()
            # Generar iniciales
            iniciales = "".join(
                [nombre[0] for nombre in nombres.split() if nombre]
            ).upper()
            ruc = data.get("ruc", "")
            return nombres, iniciales, ruc
        else:
            st.error("Error al obtener datos de SUNAT. Verifica el DNI ingresado.")
            return None, None, None
    except Exception as e:
        st.error(f"Error al conectar con la API de SUNAT: {e}")
        return None, None, None

# Función para generar el gráfico de barras
def generar_grafico(df_grafico, output_dir):
    plt.figure(figsize=(14, 8))  # Tamaño ajustado para mejor visualización
    bar_width = 0.15
    days = df_grafico["Día"]
    indices = range(len(days))
    num_months = len(df_grafico.columns) - 1  # Restando la columna 'Día'
    colores = plt.cm.tab20.colors  # Usar una paleta de colores diferente

    # Crear barras para cada mes
    for i, mes in enumerate(df_grafico.columns[1:]):
        barras = plt.bar(
            [x + bar_width * i for x in indices],
            df_grafico[mes],
            width=bar_width,
            label=mes,
            color=colores[i % len(colores)],
        )
        # Añadir etiquetas de datos
        for barra in barras:
            altura = barra.get_height()
            plt.text(
                barra.get_x() + barra.get_width() / 2.0,
                altura + 0.05,
                f"{altura}",
                ha="center",
                va="bottom",
                fontsize=8,
            )

    plt.xlabel("Día", fontsize=14)
    plt.ylabel("Total KG", fontsize=14)
    plt.title("Descargas Diarias por Mes", fontsize=16)
    plt.xticks(
        [x + bar_width * (num_months - 1) / 2 for x in indices], days, rotation=45
    )
    plt.legend(title="Meses")
    plt.tight_layout()

    # Guardar el gráfico
    grafico_path = os.path.join(output_dir, "grafico_barras.png")
    plt.savefig(grafico_path, dpi=600)
    plt.close()
    return grafico_path

# Función para convertir Word a PDF utilizando win32com.client con manejo de COM
def convertir_a_pdf(archivo_docx, archivo_pdf):
    try:
        logging.info(f"Convirtiendo {archivo_docx} a {archivo_pdf}")
        if not os.path.exists(archivo_docx):
            st.error(f"El archivo {archivo_docx} no existe y no se puede convertir a PDF.")
            logging.error(f"El archivo {archivo_docx} no existe.")
            return

        if platform.system() == 'Windows':
            with com_handler():
                word = win32.Dispatch('Word.Application')
                word.Visible = False
                word.DisplayAlerts = 0
                doc = word.Documents.Open(os.path.abspath(archivo_docx))
                time.sleep(2)
                doc.SaveAs(os.path.abspath(archivo_pdf), FileFormat=17)
                time.sleep(2)
                doc.Close()
                word.Quit()
        else:
            # Usar LibreOffice para la conversión en sistemas Linux
            try:
                # Intentar usar libreoffice
                subprocess.run([
                    'libreoffice', 
                    '--headless', 
                    '--convert-to', 
                    'pdf', 
                    archivo_docx,
                    '--outdir', 
                    os.path.dirname(archivo_pdf)
                ], check=True)
                
                # Renombrar el archivo si es necesario
                pdf_generado = os.path.splitext(archivo_docx)[0] + '.pdf'
                if pdf_generado != archivo_pdf:
                    os.rename(pdf_generado, archivo_pdf)
                    
            except subprocess.CalledProcessError:
                st.error("Error al convertir el documento. Asegúrate de que LibreOffice esté instalado.")
                return
            except Exception as e:
                st.error(f"Error inesperado durante la conversión: {str(e)}")
                return

        st.success(f"Documento convertido a PDF: {archivo_pdf}")
    except Exception as e:
        st.error(f"Error al convertir a PDF: {e}")
        logging.error(f"Error al convertir a PDF: {e}", exc_info=True)
        
# Función para procesar y convertir documentos adicionales a PDF
def procesar_documentos_adicionales(archivos_docx, carpeta_salida, valores, base_dir):
    try:
        for docx_name in archivos_docx:
            # Ruta correcta a la plantilla
            archivo_entrada = os.path.join(base_dir, docx_name)
            archivo_salida_docx = docx_name  # Mismo nombre para salida
            archivo_salida_pdf = docx_name.replace(".docx", ".pdf")
            archivo_salida_docx_path = os.path.join(carpeta_salida, archivo_salida_docx)
            archivo_salida_pdf_path = os.path.join(carpeta_salida, archivo_salida_pdf)

            # Verificar si el archivo existe
            if not os.path.exists(archivo_entrada):
                st.warning(f"El archivo {archivo_entrada} no se encontró.")
                continue

            # Llenar la plantilla de Word
            llenar_plantilla_word(
                archivo_entrada,
                archivo_salida_docx_path,
                valores,
                grafico_path=None,  # No se inserta gráfico aquí
            )

            # Convertir a PDF
            convertir_a_pdf(archivo_salida_docx_path, archivo_salida_pdf_path)

        st.success(
            "Documentos adicionales procesados y convertidos a PDF correctamente."
        )
    except Exception as e:
        st.error(f"Error al procesar documentos adicionales: {e}")

# Función para contar archivos en una lista
def contar_archivos(excel_files):
    try:
        return str(len(excel_files))
    except Exception as e:
        st.error(f"Error al contar los archivos: {e}")
        return "0"

# Función para crear un archivo ZIP de los documentos generados
def crear_zip_archivos(archivos, carpeta_salida):
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as zipf:
        for archivo in archivos:
            ruta_archivo = os.path.join(carpeta_salida, archivo)
            if os.path.exists(ruta_archivo):
                zipf.write(ruta_archivo, arcname=archivo)
    zip_buffer.seek(0)
    return zip_buffer

# Función para generar recomendaciones por IA
rapidapi_key = st.secrets["RAPIDAPI"]["key"]
def generar_respuesta_rapidapi(prompt_text):
    url = "https://cheapest-gpt-4-turbo-gpt-4-vision-chatgpt-openai-ai-api.p.rapidapi.com/v1/chat/completions"

    payload = {
        "messages": [{"role": "user", "content": prompt_text}],
        "model": "gpt-4",
        "max_tokens": 1000,
        "temperature": 0.7,
    }
    headers = {
        "x-rapidapi-key": rapidapi_key,
        "x-rapidapi-host": "cheapest-gpt-4-turbo-gpt-4-vision-chatgpt-openai-ai-api.p.rapidapi.com",
        "Content-Type": "application/json",
    }

    try:
        response = requests.post(url, json=payload, headers=headers)
        response.raise_for_status()  # Raise an exception for bad status codes
        result = response.json()
        return result["choices"][0]["message"]["content"].strip()
    except requests.RequestException as e:
        st.error(f"Error al comunicarse con la API de RapidAPI: {e}")
        return ""
    except Exception as e:
        st.error(f"Error inesperado: {e}")
        return ""

# Función para extraer el nombre del servicio desde el PDF
def extraer_nombre_servicio(pdf_file):
    texto_completo = ""
    try:
        with pdfplumber.open(BytesIO(pdf_file.read())) as pdf:
            for pagina in pdf.pages:
                texto_completo += pagina.extract_text()

        # Unificar el texto y eliminar saltos de línea excesivos
        texto_unido = ' '.join(texto_completo.split())

        # Patrón para encontrar el texto entre "2. OBJETO DE LA CONTRATACION" y "3. FINALIDAD PUBLICA"
        patron = r'2\.\s*OBJETO\s*DE\s*LA\s*CONTRATACION\s*(.*?)\s*3\.\s*FINALIDAD\s*PUBLICA'

        # Buscar el patrón en el texto unificado
        match = re.search(patron, texto_unido, re.DOTALL | re.IGNORECASE)

        if match:
            # Eliminar espacios adicionales y unificar el texto encontrado
            servicio = ' '.join(match.group(1).split())
            "Servicio extraído:", servicio # Depuración
            return servicio
        st.warning("Servicio no encontrado en el PDF.")
        return "Servicio no encontrado"
    except Exception as e:
        st.error(f"Error al extraer el servicio del PDF: {e}")
        st.exception(e)
        return "Servicio no encontrado"

# Función para extraer las actividades desde el PDF
def extraer_actividades(pdf_file):
    texto_completo = ""
    try:
        # Reiniciar el buffer del archivo PDF
        pdf_file.seek(0)
        with pdfplumber.open(BytesIO(pdf_file.read())) as pdf:
            for pagina in pdf.pages:
                texto_completo += pagina.extract_text()

        # Unificar el texto y eliminar saltos de línea excesivos
        texto_unido = ' '.join(texto_completo.split())

        # Patrón para encontrar el texto entre "7.1 ACTIVIDADES" y "7.2 PROCESO Y METODOLOGÍA"
        patron = r'7\.1\s*ACTIVIDADES\s*(.*?)\s*7\.2\s*PROCESO\s*Y\s*METODOLOGÍA'

        # Buscar el patrón en el texto unificado
        match = re.search(patron, texto_unido, re.DOTALL | re.IGNORECASE)

        if match:
            # Extraer el texto encontrado entre las dos secciones
            actividades_texto = match.group(1)

            # Dividir el texto en base a los indicadores "a)", "b)", "c)", etc.
            actividades = re.split(r'\b([a-e])\)\s*', actividades_texto)

            # Crear variables para cada actividad
            act_a = act_b = act_c = act_d = act_e = None

            # Iterar sobre las actividades encontradas y asignarlas a las variables
            for i in range(1, len(actividades), 2):
                letra = actividades[i]
                contenido = ' '.join(actividades[i+1].split())
                if letra == 'a':
                    act_a = contenido
                elif letra == 'b':
                    act_b = contenido
                elif letra == 'c':
                    act_c = contenido
                elif letra == 'd':
                    act_d = contenido
                elif letra == 'e':
                    act_e = contenido

            st.write(f"Actividades extraídas: a): {act_a}, b): {act_b}, c): {act_c}, d): {act_d}, e): {act_e}")  # Depuración
            return act_a, act_b, act_c, act_d, act_e

        st.warning("Actividades no encontradas en el PDF.")
        return (
            "Actividades no encontradas",
            "Actividades no encontradas",
            "Actividades no encontradas",
            "Actividades no encontradas",
            "Actividades no encontradas",
        )
    except Exception as e:
        st.error(f"Error al extraer las actividades del PDF: {e}")
        st.exception(e)
        return (
            "Actividades no encontradas",
            "Actividades no encontradas",
            "Actividades no encontradas",
            "Actividades no encontradas",
            "Actividades no encontradas",
        )

# Función para reemplazar texto en párrafos y tablas de manera robusta
def reemplazar_texto(doc, valores):
    """
    Reemplaza los marcadores de posición en todo el documento, incluyendo párrafos y tablas.
    """
    # Reemplazo en párrafos
    for paragraph in doc.paragraphs:
        for key, value in valores.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                inline = paragraph.runs
                # Reemplazar en cada run
                for run in inline:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
    
    # Reemplazo en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in valores.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in paragraph.text:
                            inline = paragraph.runs
                            for run in inline:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))

def limpiar_valor(valor):
    """
    Limpia y convierte el valor según las siguientes reglas:
    - Si el valor es 'S/R' (ignorando mayúsculas/minúsculas), devuelve 0.
    - Si el valor es una cadena que contiene ',' o ';', reemplaza estos caracteres por '.' y convierte a float.
    - Si el valor es numérico, lo devuelve como float.
    - En cualquier otro caso, devuelve 0.
    """
    if isinstance(valor, str):
        if valor.strip().upper() == "S/R":
            return 0.0
        # Reemplazar ',' y ';' por '.' y tratar de convertir a float
        valor = valor.replace(",", ".").replace(";", ".")
        try:
            return float(valor)
        except ValueError:
            return 0.0
    elif isinstance(valor, (int, float)):
        return float(valor)
    else:
        return 0.0

def generar_data_excel(excel_files, output_file):
    try:
        # Inicializar listas para almacenar datos de todas las hojas
        total_data = []
        comentarios_data = []
        grafico_data_list = []
        dfs_descripcion = []
        all_precios = pd.DataFrame()

        # Diccionario para los meses en español
        meses_espanol = MESES_ES.copy()

        # Diccionario para almacenar los rangos de hora por fecha
        hora_range_data = {}

        # Bucle para procesar cada archivo Excel
        for file in excel_files:
            # Leer todas las hojas del archivo Excel
            xls = pd.ExcelFile(file, engine='openpyxl')
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name, header=None)

                # Procesamiento similar al actual pero para cada hoja
                # Asegúrate de que cada hoja tenga la estructura esperada
                date = df.iloc[7, 11]  # Fecha en L8 (0-based)
                observation = df.iloc[12, 22]  # Observación en W13 (0-based)
                total_value = df.iloc[11, 6]  # Total en G12 (0-based)

                # Obtener G13:G57 y X13:X57
                G_values = df.iloc[12:57, 6]  # G13:G57
                X_values = df.iloc[12:57, 23]  # X13:X57

                # Crear un DataFrame para Pescados y Mariscos
                total_items_df = pd.DataFrame({"G": G_values, "X": X_values})

                # Sumar los valores para Pescados y Mariscos
                sum_pescado = total_items_df.loc[
                    total_items_df["X"] == "PESCADOS", "G"
                ].sum()
                sum_marisco = total_items_df.loc[
                    total_items_df["X"] == "MARISCOS", "G"
                ].sum()

                # Manejar el valor total
                if pd.isna(total_value):
                    total_value = 0.0
                else:
                    try:
                        total_value = float(total_value)
                    except ValueError:
                        total_value = 0.0

                # Agregar los datos a la lista total_data
                total_data.append(
                    {
                        "Fecha": date,
                        "Pescado (KG)": sum_pescado,
                        "Marisco (KG)": sum_marisco,
                        "Total (KG)": total_value,
                    }
                )

                # Agregar datos para la hoja 'grafico'
                fecha_dt = (
                    pd.to_datetime(date, format="%d %m %Y", errors="coerce")
                    if pd.notna(date)
                    else pd.NaT
                )
                if pd.notna(fecha_dt):
                    dia = fecha_dt.day
                    mes_num = fecha_dt.month
                    mes_nombre = meses_espanol.get(mes_num, "Mes desconocido")
                    grafico_data_list.append(
                        {"Día": dia, "Mes": mes_nombre, "Total_Kg": total_value}
                    )

                # Obtener 'Primer Valor' y 'Último Valor' de 'Hora'
                hora_col = df.iloc[12:57, 14]
                primer_valor = df.iloc[12, 14]
                primer_valor = "" if pd.isna(primer_valor) else primer_valor

                # Encontrar el último valor no nulo en la columna 'Hora'
                ultimo_valor = (
                    hora_col.dropna().iloc[-1]
                    if not hora_col.dropna().empty
                    else primer_valor
                )

                # Convertir 'primer_valor' y 'ultimo_valor' a formato 'HH:MM'
                def format_hora(valor):
                    if isinstance(valor, pd.Timestamp):
                        return valor.strftime("%H:%M")
                    elif isinstance(valor, str):
                        try:
                            return pd.to_datetime(valor).strftime("%H:%M")
                        except ValueError:
                            return valor
                    else:
                        return valor

                primer_valor_formateado = format_hora(primer_valor)
                ultimo_valor_formateado = format_hora(ultimo_valor)

                hora_range = f"{primer_valor_formateado} - {ultimo_valor_formateado}"

                # Almacenar el 'Hora' range por 'Fecha'
                hora_range_data[date] = hora_range

                # Leer columnas adicionales de 'Tamaño' y 'Precio'
                tamaño_columns = [15, 17, 19]  # P, R, T
                tamaño_data = df.iloc[12:57, tamaño_columns]
                
                # Aplicar limpieza a las columnas de 'Tamaño'
                tamaño_data_limpio = tamaño_data.applymap(limpiar_valor)
                tamaño_combined = tamaño_data_limpio.stack().groupby(level=0).mean()

                # Convertir 'tamaño_combined' a numérico y llenar NaN con 0
                tamaño_combined = pd.to_numeric(tamaño_combined, errors="coerce").fillna(0.0)

                precio_columns = [16, 18, 20]  # Q, S, U
                precio_data = df.iloc[12:57, precio_columns]
                
                # Aplicar limpieza a las columnas de 'Precio'
                precio_data_limpio = precio_data.applymap(limpiar_valor)
                
                # Incluir 'Nombre común' antes de melt para mantener la asociación
                nombre_comun = df.iloc[
                    12:57, 2
                ].values  # Suponiendo que la columna 2 contiene 'Nombre común'
                precio_data_with_nombre = precio_data_limpio.copy()
                precio_data_with_nombre["Nombre común"] = nombre_comun

                # Realizar el melt incluyendo 'Nombre común' como id_vars
                precio_melted = precio_data_with_nombre.melt(
                    id_vars="Nombre común", value_vars=precio_columns, value_name="Precio"
                )

                # Convertir 'Precio' a numérico y filtrar valores > 0
                precio_melted["Precio"] = pd.to_numeric(
                    precio_melted["Precio"], errors="coerce"
                )
                precio_melted = precio_melted.dropna(subset=["Precio"])
                precio_melted = precio_melted[precio_melted["Precio"] > 0]

                # Agregar los precios individuales a 'all_precios'
                all_precios = pd.concat(
                    [all_precios, precio_melted[["Nombre común", "Precio"]]],
                    ignore_index=True,
                )

                # Crear DataFrame temporal para descripción
                temp_df = pd.DataFrame(
                    {
                        "Fecha": [date] * len(df.iloc[12:57, 2]),
                        "Nombre_común": df.iloc[12:57, 2],
                        "Nombre_científico": df.iloc[12:57, 3],
                        "Volumen_Kg": df.iloc[12:57, 6],
                        "Procedencia": df.iloc[12:57, 8],
                        "Aparejo": df.iloc[12:57, 7],
                        "Hora": df.iloc[12:57, 14],
                        "Tamaño": tamaño_combined,
                        "Precio": precio_data_limpio.mean(axis=1),  # Promedio para descripción
                        "Observación": "",
                    }
                )

                temp_df = temp_df.dropna(
                    subset=["Nombre_común", "Volumen_Kg", "Hora"], how="all"
                )

                if not temp_df.empty:
                    temp_df.iloc[0, temp_df.columns.get_loc("Observación")] = observation
                    dfs_descripcion.append(temp_df)
                elif pd.notna(observation):
                    # Manejar casos donde solo hay observación
                    observation_df = pd.DataFrame(
                        {
                            "Fecha": [date],
                            "Nombre_común": ["N/A"],
                            "Nombre_científico": ["N/A"],
                            "Volumen_Kg": [0.0],
                            "Procedencia": ["N/A"],
                            "Aparejo": ["N/A"],
                            "Hora": ["N/A"],
                            "Tamaño": [0.0],
                            "Precio": [0.0],
                            "Observación": [observation],
                        }
                    )
                    dfs_descripcion.append(observation_df)
                    hora_range_data[date] = ""  # Sin rango de hora

        # Continuar con el procesamiento como ya lo tienes
        if not dfs_descripcion:
            st.error("No se encontraron datos de descripción para procesar.")
            return

        # Combinar todos los DataFrames de descripción
        result_df = pd.concat(dfs_descripcion, ignore_index=True)
        result_df["Fecha"] = pd.to_datetime(
            result_df["Fecha"], format="%d %m %Y", errors="coerce"
        )
        result_df = result_df.sort_values("Fecha")

        # Agrupar y resumir datos para descripción
        grouped_df = (
            result_df.groupby(["Fecha", "Nombre_común"])
            .agg(
                {
                    "Volumen_Kg": "sum",
                    "Tamaño": "mean",
                    "Precio": "mean",  # Usamos el promedio para descripción
                    "Observación": "first",
                    "Nombre_científico": "first",
                }
            )
            .reset_index()
        )

        grouped_df["Tamaño"] = grouped_df["Tamaño"].round(2)
        grouped_df["Precio"] = grouped_df["Precio"].round(2)

        # Limpiar nombres de columnas para evitar espacios adicionales
        grouped_df.columns = grouped_df.columns.str.strip()

        # Formatear 'Fecha' para merge
        grouped_df["Fecha_str"] = grouped_df["Fecha"].dt.strftime("%d %m %Y")

        # Crear DataFrame de 'hora_range_data'
        hora_range_df = pd.DataFrame(
            list(hora_range_data.items()), columns=["Fecha_str", "Hora_Range"]
        )

        # Unir 'hora_range_df' con 'grouped_df'
        grouped_df = grouped_df.merge(hora_range_df, on="Fecha_str", how="left")
        grouped_df["Hora"] = grouped_df["Hora_Range"]
        grouped_df = grouped_df.drop(["Hora_Range", "Fecha_str"], axis=1)

        # Formatear 'Fecha' para visualización
        grouped_df["Fecha"] = grouped_df["Fecha"].dt.strftime("%d/%m/%Y")

        # Crear la columna 'Descripcion' antes de seleccionar las columnas
        def create_description(row):
            if (
                row["Nombre_común"] == "N/A"
                and row["Volumen_Kg"] == 0.0
                and row["Tamaño"] == 0.0
                and row["Precio"] == 0.0
            ):
                return row["Observación"]
            else:
                precio_formateado = "{:.2f}".format(row["Precio"])
                return f"{row['Nombre_común']} con {row['Volumen_Kg']} kg; talla promedio de {row['Tamaño']} cm; precio por kilo de S/ {precio_formateado}"

        grouped_df["Descripcion"] = grouped_df.apply(create_description, axis=1)

        # Verificar que la columna 'Descripcion' se ha creado correctamente
        if "Descripcion" not in grouped_df.columns:
            raise KeyError("La columna 'Descripcion' no se ha creado correctamente.")

        # Definir el orden de las columnas según lo solicitado
        columns_order = [
            "Fecha",
            "Nombre_común",
            "Volumen_Kg",
            "Hora",
            "Tamaño",
            "Precio",
            "Observación",
            "Descripcion",
        ]

        # Asegurarse de que todas las columnas existan antes de seleccionarlas
        missing_columns = [
            col for col in columns_order if col not in grouped_df.columns
        ]
        if missing_columns:
            raise KeyError(
                f"Las siguientes columnas faltan en 'grouped_df': {missing_columns}"
            )

        # Reordenar columnas
        grouped_df = grouped_df[columns_order]

        # Ordenar por 'Fecha' y 'Nombre_común' para asegurar el orden correcto
        grouped_df = grouped_df.sort_values(["Fecha", "Nombre_común"])

        # Establecer 'Hora' a cadena vacía excepto para la primera ocurrencia de cada fecha
        grouped_df.loc[grouped_df.duplicated(subset="Fecha"), "Hora"] = ""

        # Crear 'total_df' con las nuevas columnas
        total_df = pd.DataFrame(total_data)
        total_df["Fecha"] = pd.to_datetime(
            total_df["Fecha"], format="%d %m %Y", errors="coerce"
        )
        total_df = total_df.sort_values("Fecha")

        # Convertir las columnas a numérico
        total_df["Pescado (KG)"] = pd.to_numeric(
            total_df["Pescado (KG)"], errors="coerce"
        ).fillna(0.0)
        total_df["Marisco (KG)"] = pd.to_numeric(
            total_df["Marisco (KG)"], errors="coerce"
        ).fillna(0.0)
        total_df["Total (KG)"] = pd.to_numeric(
            total_df["Total (KG)"], errors="coerce"
        ).fillna(0.0)

        # Formatear 'Fecha' para visualización
        total_df["Fecha"] = total_df["Fecha"].dt.strftime("%d/%m/%Y")

        # Calcular los totales
        total_sum_pescado = total_df["Pescado (KG)"].sum()
        total_sum_marisco = total_df["Marisco (KG)"].sum()
        total_sum_total = total_df["Total (KG)"].sum()

        # Crear la fila total
        total_row = pd.DataFrame(
            {
                "Fecha": ["TOTAL"],
                "Pescado (KG)": [total_sum_pescado],
                "Marisco (KG)": [total_sum_marisco],
                "Total (KG)": [total_sum_total],
            }
        )

        # Concatenar la fila total al DataFrame total_df
        total_df = pd.concat([total_df, total_row], ignore_index=True)

        # Reemplazar 0 con "–" en las columnas numéricas excepto en la fila 'TOTAL'
        numeric_columns = ["Pescado (KG)", "Marisco (KG)", "Total (KG)"]
        for col in numeric_columns:
            # Convertir a object para permitir strings
            total_df[col] = total_df[col].astype(object)
            # Crear máscara para filas que no sean 'TOTAL' y donde el valor sea 0
            mask = (total_df["Fecha"] != "TOTAL") & (total_df[col] == 0.0)
            # Asignar "–" donde la máscara es True
            total_df.loc[mask, col] = "–"

        # Filtrar registros válidos (excluir N/A)
        valid_df = result_df[
            (result_df["Nombre_común"] != "N/A") & (result_df["Nombre_común"].notna())
        ]

        # Calcular el total de volumen por especie
        volumen_df = (
            valid_df.groupby("Nombre_común")
            .agg(KG_POR_ESPECIE=("Volumen_Kg", "sum"))
            .reset_index()
        )

        # Agrupar y calcular precios mínimos y máximos a partir de 'all_precios'
        if not all_precios.empty:
            # Filtrar precios > 0 para PRECIO MAS BAJO
            precios_mas_bajo_df = (
                all_precios[all_precios["Precio"] > 0]
                .groupby("Nombre común")
                .agg(PRECIO_MAS_BAJO=("Precio", "min"))
                .reset_index()
            )

            # Obtener PRECIO MAS ALTO
            precios_mas_alto_df = (
                all_precios.groupby("Nombre común")
                .agg(PRECIO_MAS_ALTO=("Precio", "max"))
                .reset_index()
            )

            # Combinar ambos DataFrames
            precios_df = pd.merge(
                precios_mas_bajo_df, precios_mas_alto_df, on="Nombre común", how="outer"
            )
        else:
            precios_df = pd.DataFrame(
                columns=["Nombre común", "PRECIO MAS BAJO (S/)", "PRECIO MAS ALTO (S/)"]
            )

        # Renombrar 'Nombre común' a 'Nombre_común' para consistencia
        precios_df.rename(columns={"Nombre común": "Nombre_común"}, inplace=True)

        # Unir todos los DataFrames
        especies_df = volumen_df.merge(precios_df, on="Nombre_común", how="left")

        # Agregar 'NOMBRE CIENTIFICO'
        nombres_cientificos = valid_df[
            ["Nombre_común", "Nombre_científico"]
        ].drop_duplicates()
        especies_df = especies_df.merge(
            nombres_cientificos, on="Nombre_común", how="left"
        )

        # Renombrar columnas correctamente
        especies_df.rename(
            columns={
                "Nombre_común": "NOMBRE COMUN",
                "Nombre_científico": "NOMBRE CIENTIFICO",
                "PRECIO_MAS_BAJO": "PRECIO MAS BAJO (S/)",
                "PRECIO_MAS_ALTO": "PRECIO MAS ALTO (S/)",
                "KG_POR_ESPECIE": "KG POR ESPECIE",
            },
            inplace=True,
        )

        # Limpiar nombres de columnas para evitar espacios adicionales
        especies_df.columns = especies_df.columns.str.strip()

        # Manejar valores NaN en precios
        if "PRECIO MAS BAJO (S/)" in especies_df.columns:
            especies_df["PRECIO MAS BAJO (S/)"] = especies_df[
                "PRECIO MAS BAJO (S/)"
            ].fillna("–")
        else:
            especies_df["PRECIO MAS BAJO (S/)"] = "–"

        if "PRECIO MAS ALTO (S/)" in especies_df.columns:
            especies_df["PRECIO MAS ALTO (S/)"] = especies_df[
                "PRECIO MAS ALTO (S/)"
            ].fillna("–")
        else:
            especies_df["PRECIO MAS ALTO (S/)"] = "–"

        # Definir el orden de las columnas según lo solicitado
        especies_columns_order = [
            "NOMBRE COMUN",
            "NOMBRE CIENTIFICO",
            "PRECIO MAS BAJO (S/)",
            "PRECIO MAS ALTO (S/)",
            "KG POR ESPECIE",
        ]

        # Reordenar las columnas
        especies_df = especies_df[especies_columns_order]

        # Ordenar el DataFrame
        especies_df = especies_df.sort_values("NOMBRE COMUN")

        # Crear 'procedencia_df' excluyendo 'N/A' y valores nulos
        procedencia_df = result_df.dropna(subset=["Procedencia"])
        procedencia_df = procedencia_df[procedencia_df["Procedencia"] != "N/A"]
        procedencia_df = (
            procedencia_df.groupby("Procedencia")
            .agg(Volumen_Kg=("Volumen_Kg", "sum"))
            .reset_index()
        )
        procedencia_df = procedencia_df.sort_values("Procedencia")

        # Crear 'aparejos_data' agrupando 'result_df' por 'Aparejo'
        aparejos_df = (
            result_df.groupby("Aparejo")
            .agg(Volumen_Kg=("Volumen_Kg", "sum"))
            .reset_index()
        )
        aparejos_df = aparejos_df.sort_values("Aparejo")

        # Convertir 'aparejos_df' a diccionario
        aparejos_data = dict(zip(aparejos_df["Aparejo"], aparejos_df["Volumen_Kg"]))

        # Crear 'comentarios_df'
        # Especie principal
        if not especies_df.empty:
            especie_principal = especies_df.loc[especies_df["KG POR ESPECIE"].idxmax()]
            comentarios_data.extend(
                [
                    {
                        "Tipo_comentario": "Especie_principal_1",
                        "Comentario": f"{especie_principal['NOMBRE COMUN']}",
                    },
                    {
                        "Tipo_comentario": "Especie_principal_2",
                        "Comentario": f"({especie_principal['NOMBRE CIENTIFICO']})",
                    },
                    {
                        "Tipo_comentario": "Especie_principal_3",
                        "Comentario": f"{especie_principal['KG POR ESPECIE']}",
                    },
                ]
            )
        else:
            comentarios_data.extend(
                [
                    {
                        "Tipo_comentario": "Especie_principal_1",
                        "Comentario": "No hay datos disponibles",
                    },
                    {"Tipo_comentario": "Especie_principal_2", "Comentario": ""},
                    {"Tipo_comentario": "Especie_principal_3", "Comentario": ""},
                ]
            )

        # Especie mínima (excluyendo KG POR ESPECIE igual a cero)
        especies_df_nonzero = especies_df[especies_df["KG POR ESPECIE"] > 0]
        if not especies_df_nonzero.empty:
            especie_minima = especies_df_nonzero.loc[
                especies_df_nonzero["KG POR ESPECIE"].idxmin()
            ]
            comentarios_data.extend(
                [
                    {
                        "Tipo_comentario": "Especie_minima_1",
                        "Comentario": f"{especie_minima['NOMBRE COMUN']}",
                    },
                    {
                        "Tipo_comentario": "Especie_minima_2",
                        "Comentario": f"({especie_minima['NOMBRE CIENTIFICO']})",
                    },
                    {
                        "Tipo_comentario": "Especie_minima_3",
                        "Comentario": f"{especie_minima['KG POR ESPECIE']}",
                    },
                ]
            )
        else:
            comentarios_data.extend(
                [
                    {
                        "Tipo_comentario": "Especie_minima_1",
                        "Comentario": "No hay datos disponibles",
                    },
                    {"Tipo_comentario": "Especie_minima_2", "Comentario": ""},
                    {"Tipo_comentario": "Especie_minima_3", "Comentario": ""},
                ]
            )

        # Procesar datos para mayor y menor día
        total_df_numeric = total_df.iloc[:-1].copy()  # Excluir la última fila 'TOTAL'
        total_df_numeric["Total (KG)"] = pd.to_numeric(
            total_df_numeric["Total (KG)"], errors="coerce"
        )

        if not total_df_numeric.empty:
            # Encontrar máximo total
            max_total = total_df_numeric["Total (KG)"].max()
            # Encontrar todas las fechas donde 'Total (KG)' es igual a max_total
            max_total_dates_df = total_df_numeric[
                total_df_numeric["Total (KG)"] == max_total
            ]
            max_total_dates_list = pd.to_datetime(
                max_total_dates_df["Fecha"], format="%d/%m/%Y", errors="coerce"
            ).dt.date.tolist()
            # Procesar fechas para crear 'mayor_dia_fecha'
            dates_info = []
            for date in max_total_dates_list:
                dia = date.day
                mes_num = date.month
                mes_nombre = meses_espanol.get(mes_num, "Mes desconocido")
                dates_info.append(
                    {"dia": dia, "mes_num": mes_num, "mes_nombre": mes_nombre}
                )

            # Agrupar fechas por mes
            dates_by_month = defaultdict(list)
            for info in dates_info:
                dates_by_month[info["mes_num"]].append(info)

            # Construir las cadenas de fechas
            date_strings = []
            for mes_num, infos in dates_by_month.items():
                dias = [str(info["dia"]) for info in infos]
                mes_nombre = infos[0]["mes_nombre"]
                if len(dias) > 1:
                    dias_str = ", ".join(dias[:-1]) + " y " + dias[-1]
                else:
                    dias_str = dias[0]
                date_strings.append(f"{dias_str} de {mes_nombre}")

            mayor_dia_fecha = "; ".join(date_strings)

            comentarios_data.append(
                {
                    "Tipo_comentario": "Mayor_dia",
                    "Comentario": f"{mayor_dia_fecha} se presentó la mayor descarga con un total de {max_total} kg",
                }
            )

            # Para 'Menor_dia' (excluyendo ceros)
            menor_dia_df = total_df_numeric[total_df_numeric["Total (KG)"] > 0]
            if not menor_dia_df.empty:
                min_total = menor_dia_df["Total (KG)"].min()
                # Encontrar todas las fechas donde 'Total (KG)' es igual a min_total
                min_total_dates_df = menor_dia_df[
                    menor_dia_df["Total (KG)"] == min_total
                ]
                min_total_dates_list = pd.to_datetime(
                    min_total_dates_df["Fecha"], format="%d/%m/%Y", errors="coerce"
                ).dt.date.tolist()
                # Procesar fechas para crear 'menor_dia_fecha'
                dates_info = []
                for date in min_total_dates_list:
                    dia = date.day
                    mes_num = date.month
                    mes_nombre = meses_espanol.get(mes_num, "Mes desconocido")
                    dates_info.append(
                        {"dia": dia, "mes_num": mes_num, "mes_nombre": mes_nombre}
                    )

                # Agrupar fechas por mes
                dates_by_month = defaultdict(list)
                for info in dates_info:
                    dates_by_month[info["mes_num"]].append(info)

                # Construir las cadenas de fechas
                date_strings = []
                for mes_num, infos in dates_by_month.items():
                    dias = [str(info["dia"]) for info in infos]
                    mes_nombre = infos[0]["mes_nombre"]
                    if len(dias) > 1:
                        dias_str = ", ".join(dias[:-1]) + " y " + dias[-1]
                    else:
                        dias_str = dias[0]
                    date_strings.append(f"{dias_str} de {mes_nombre}")

                menor_dia_fecha = "; ".join(date_strings)

                comentarios_data.append(
                    {
                        "Tipo_comentario": "Menor_dia",
                        "Comentario": f"{menor_dia_fecha} con {min_total} kg",
                    }
                )
            else:
                comentarios_data.append(
                    {
                        "Tipo_comentario": "Menor_dia",
                        "Comentario": "No se encontró un día con valor mayor a 0 kg",
                    }
                )
        else:
            comentarios_data.extend(
                [
                    {
                        "Tipo_comentario": "Mayor_dia",
                        "Comentario": "No hay datos disponibles para calcular el mayor día",
                    },
                    {
                        "Tipo_comentario": "Menor_dia",
                        "Comentario": "No hay datos disponibles para calcular el menor día",
                    },
                ]
            )

        # Aparejos
        aparejos_sorted = sorted(
            aparejos_data.items(), key=lambda x: x[1], reverse=True
        )
        aparejos_comment = " y ".join(
            [
                f"{aparejo} con {valor} Kg"
                for aparejo, valor in aparejos_sorted
                if valor > 0
            ]
        )
        comentarios_data.append(
            {"Tipo_comentario": "Aparejos", "Comentario": aparejos_comment}
        )

        # Para los comentarios de precios usando especies_df
        if not especies_df.empty:
            # Asegurarse de que los precios sean de tipo float
            especies_df["PRECIO MAS BAJO (S/)"] = pd.to_numeric(
                especies_df["PRECIO MAS BAJO (S/)"], errors="coerce"
            )
            especies_df["PRECIO MAS ALTO (S/)"] = pd.to_numeric(
                especies_df["PRECIO MAS ALTO (S/)"], errors="coerce"
            )

            # Encontrar el precio más alto entre todos los registros con precio > 0
            max_precio = especies_df["PRECIO MAS ALTO (S/)"].max()
            max_precio_registros = especies_df[
                especies_df["PRECIO MAS ALTO (S/)"] == max_precio
            ]

            for _, row in max_precio_registros.iterrows():
                comentarios_data.extend(
                    [
                        {
                            "Tipo_comentario": "Mayor_Precio_1",
                            "Comentario": f"{row['NOMBRE COMUN']}",
                        },
                        {
                            "Tipo_comentario": "Mayor_Precio_2",
                            "Comentario": f"({row['NOMBRE CIENTIFICO']})",
                        },
                        {
                            "Tipo_comentario": "Mayor_Precio_3",
                            "Comentario": f"{row['PRECIO MAS ALTO (S/)']}",
                        },
                    ]
                )

            # Encontrar el precio más bajo entre todos los registros con precio > 0
            min_precio = especies_df["PRECIO MAS BAJO (S/)"].min()
            min_precio_registros = especies_df[
                especies_df["PRECIO MAS BAJO (S/)"] == min_precio
            ]

            for _, row in min_precio_registros.iterrows():
                comentarios_data.extend(
                    [
                        {
                            "Tipo_comentario": "Menor_Precio_1",
                            "Comentario": f"{row['NOMBRE COMUN']}",
                        },
                        {
                            "Tipo_comentario": "Menor_Precio_2",
                            "Comentario": f"({row['NOMBRE CIENTIFICO']})",
                        },
                        {
                            "Tipo_comentario": "Menor_Precio_3",
                            "Comentario": f"{row['PRECIO MAS BAJO (S/)']}",
                        },
                    ]
                )
        else:
            comentarios_data.extend(
                [
                    {
                        "Tipo_comentario": "Mayor_Precio_1",
                        "Comentario": "No hay datos disponibles",
                    },
                    {"Tipo_comentario": "Mayor_Precio_2", "Comentario": ""},
                    {"Tipo_comentario": "Mayor_Precio_3", "Comentario": ""},
                    {
                        "Tipo_comentario": "Menor_Precio_1",
                        "Comentario": "No hay datos disponibles",
                    },
                    {"Tipo_comentario": "Menor_Precio_2", "Comentario": ""},
                    {"Tipo_comentario": "Menor_Precio_3", "Comentario": ""},
                ]
            )

        # Total de G12
        total_sum_G12 = sum(
            [
                d["Total (KG)"] if isinstance(d["Total (KG)"], (int, float)) else 0
                for d in total_data
            ]
        )
        comentarios_data.append(
            {"Tipo_comentario": "Total", "Comentario": f"{total_sum_G12} kg"}
        )

        # Crear DataFrame de comentarios
        comentarios_df = pd.DataFrame(comentarios_data)

        # Preparar datos para la hoja 'grafico'
        grafico_df = pd.DataFrame(grafico_data_list)

        # Crear pivot table sin reindexar a todos los días
        grafico_pivot = grafico_df.pivot_table(
            index="Día", columns="Mes", values="Total_Kg", aggfunc="sum", fill_value=0
        ).reset_index()

        # Renombrar columnas de meses
        grafico_pivot.columns = ["Día"] + [
            meses_espanol.get(mes_num, f"{mes_num}")
            for mes_num in grafico_pivot.columns[1:]
        ]

        # Ordenar por 'Día'
        grafico_pivot = grafico_pivot.sort_values("Día")

        # Convertir todos los valores a int para evitar problemas de formato en Excel
        grafico_pivot.iloc[:, 1:] = grafico_pivot.iloc[:, 1:].astype(int)

        # Guardar los DataFrames en el archivo Excel
        with pd.ExcelWriter(output_file, engine="openpyxl") as writer:
            # Hoja 'descripcion'
            grouped_df.to_excel(writer, index=False, sheet_name="descripcion")

            # Hoja 'total'
            total_df.to_excel(writer, index=False, sheet_name="total")

            # Hoja 'especies'
            especies_df.to_excel(writer, index=False, sheet_name="especies")

            # Hoja 'procedencia'
            procedencia_df.to_excel(writer, index=False, sheet_name="procedencia")

            # Hoja 'comentarios'
            comentarios_df.to_excel(writer, index=False, sheet_name="comentarios")

            # Hoja 'grafico'
            grafico_pivot.to_excel(writer, index=False, sheet_name="grafico")

            # Obtener el libro de trabajo
            workbook = writer.book

            # Formato para la hoja 'especies'
            if "especies" in workbook.sheetnames:
                worksheet = workbook["especies"]
                for row in worksheet.iter_rows(
                    min_row=2, min_col=2, max_col=2, max_row=worksheet.max_row
                ):
                    for cell in row:
                        cell.font = Font(italic=True)
                for col in ["A", "C", "D", "E"]:
                    for cell in worksheet[col + "1:" + col + str(worksheet.max_row)]:
                        cell[0].font = Font(bold=True)

            # Formato para la columna 'Fecha' en 'total'
            if "total" in workbook.sheetnames:
                worksheet = workbook["total"]
                for row in worksheet.iter_rows(
                    min_row=2, min_col=1, max_col=1, max_row=worksheet.max_row
                ):
                    for cell in row:
                        if cell.value != "TOTAL":
                            cell.number_format = "DD/MM/YYYY"

            # Formato para la hoja 'comentarios'
            if "comentarios" in workbook.sheetnames:
                worksheet = workbook["comentarios"]
                italic_font = Font(italic=True)
                for row in worksheet.iter_rows(
                    min_row=2, max_row=worksheet.max_row, min_col=1, max_col=2
                ):
                    if row[0].value in [
                        "Especie_principal_2",
                        "Especie_minima_2",
                        "Mayor_Precio_2",
                        "Menor_Precio_2",
                    ]:
                        row[1].font = italic_font

            # Formato para la hoja 'grafico'
            if "grafico" in workbook.sheetnames:
                worksheet = workbook["grafico"]
                # Ajustar el ancho de las columnas para mejor legibilidad
                for column_cells in worksheet.columns:
                    length = max(len(str(cell.value)) for cell in column_cells)
                    adjusted_width = length + 2
                    worksheet.column_dimensions[column_cells[0].column_letter].width = (
                        adjusted_width
                    )

        st.success(f"Archivo generado: {output_file}")
        return True
    except Exception as e:
        st.error(f"Error al analizar los datos: {e}")
        return False

def llenar_primera_tabla(doc, df_descripcion):
    if df_descripcion is None or df_descripcion.empty:
        logging.error("Línea 1225: El DataFrame 'df_descripcion' es None o está vacío.")
        return

    try:
        table = doc.tables[0]  # Primera tabla

        # Limpiar filas existentes excepto el encabezado
        for row in table.rows[1:]:
            tbl = table._tbl
            tbl.remove(row._tr)

        # Agrupar por fecha y concatenar descripciones
        df_grouped = df_descripcion.groupby('Fecha').agg({
            'Hora': 'first',  # Tomar la primera hora de cada fecha
            'Descripcion': lambda x: '\n'.join(x)  # Concatenar descripciones
        }).reset_index()

        # Convertir fecha al formato correcto
        df_grouped['Fecha'] = pd.to_datetime(df_grouped['Fecha'], format='%d/%m/%Y').dt.strftime('%d/%m/%Y')

        # Llenar la tabla
        for _, row in df_grouped.iterrows():
            new_row = table.add_row().cells
            new_row[0].text = row['Fecha']
            new_row[1].text = str(row['Hora']) if pd.notnull(row['Hora']) else ''
            new_row[2].text = str(row['Descripcion'])

            # Aplicar formato
            for cell in new_row:
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    except Exception as e:
        logging.error(f"Error al llenar la primera tabla: {e}")

def llenar_segunda_tabla(doc, df_total):
    if df_total is None or df_total.empty:
        logging.error("Linea 1322: El DataFrame 'df_total' es None o está vacío.")
        return

    try:
        table = doc.tables[1]  # Segunda tabla

        # Limpiar filas existentes excepto el encabezado
        for row in table.rows[1:]:
            tbl = table._tbl
            tbl.remove(row._tr)

        # Llenar la tabla
        for _, row in df_total.iterrows():
            new_row = table.add_row().cells
            new_row[0].text = str(row['Fecha'])
            new_row[1].text = str(row['Pescado (KG)'])
            new_row[2].text = str(row['Marisco (KG)'])
            new_row[3].text = str(row['Total (KG)'])

            # Centrar valores
            for cell in new_row:
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    except Exception as e:
        logging.error(f"Error al llenar la segunda tabla: {e}")

def llenar_tercera_tabla(doc, df_especies):
    try:
        table = doc.tables[2]  # Tercera tabla
        
        # Limpiar filas existentes excepto el encabezado
        for row in table.rows[1:]:
            tbl = table._tbl
            tbl.remove(row._tr)

        # Seleccionar solo las columnas necesarias
        columnas = ['NOMBRE COMUN', 'NOMBRE CIENTIFICO', 'PRECIO MAS BAJO (S/)', 'PRECIO MAS ALTO (S/)']
        df_especies = df_especies[columnas]

        # Llenar la tabla
        for _, row in df_especies.iterrows():
            new_row = table.add_row().cells
            for i, value in enumerate(row):
                new_row[i].text = str(value)
                
                # Aplicar formato
                paragraph = new_row[i].paragraphs[0]
                if i < 2:  # Nombre común y científico alineados a la izquierda
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:  # Precios centrados
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    except Exception as e:
        logging.error(f"Error al llenar la tercera tabla: {e}")

def llenar_tablas_procedencia(doc, df_procedencia):
    try:
        # Llenar cuarta y quinta tabla (sin encabezado)
        for table_index in [3, 4]:
            table = doc.tables[table_index]
            
            # Limpiar todas las filas (no hay encabezado)
            for row in table.rows[:]:
                tbl = table._tbl
                tbl.remove(row._tr)

            # Llenar la tabla desde la primera fila
            for _, row in df_procedencia.iterrows():
                new_row = table.add_row().cells
                new_row[0].text = str(row['Procedencia'])
                new_row[1].text = str(row['Volumen_Kg'])

                # Alinear texto
                new_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                new_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    except Exception as e:
        logging.error(f"Error al llenar las tablas de procedencia: {e}")

def add_species_images_and_captions(doc, uploaded_images, imagenes_dir, especies_df):
    if not uploaded_images:
        logging.warning("No hay imágenes para insertar.")
        return
    # Add a new paragraph for the title
    doc.add_paragraph("FOTOGRAFÍAS").style = 'Heading 1'

    fig_num = 2  # Iniciar numeración de figuras

    for image in uploaded_images:
        nombre_comun = os.path.splitext(image.name)[0].upper()
        imagen_path = os.path.join(imagenes_dir, image.name)
        
        # Buscar los datos correspondientes en especies_df
        especie_data = especies_df[especies_df['NOMBRE COMUN'] == nombre_comun]
        
        if not especie_data.empty:
            nombre_cientifico = especie_data['NOMBRE CIENTIFICO'].iloc[0]
            kg_especie = especie_data['KG POR ESPECIE'].iloc[0]
        else:
            nombre_cientifico = "Nombre científico no encontrado"
            kg_especie = 0
            
        if not os.path.exists(imagen_path):
            st.error(f"La imagen {image.name} no se encontró en {imagenes_dir}.")
            logging.error(f"La imagen {image.name} no se encontró en {imagenes_dir}.")
        else:
            logging.info(f"Imagen {image.name} encontrada en {imagenes_dir}.")

        if os.path.exists(imagen_path):
            # Añadir la imagen
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(imagen_path, width=Inches(4.331))

            # Añadir el pie de foto con el formato especificado
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Primera parte del pie de foto (Fig. XX y nombre común)
            caption_run = caption.add_run(f"Fig. {fig_num:02d} Ejemplares de {nombre_comun} (")
            caption_run.italic = False
            
            # Nombre científico en cursiva
            scientific_name_run = caption.add_run(f"{nombre_cientifico}")
            scientific_name_run.italic = True
            
            # Parte final con los kg
            kg_run = caption.add_run(f") con una descarga de {kg_especie} Kg")
            kg_run.italic = False

            fig_num += 1

def obtener_comentarios_excel(comentarios_df):
    comentarios = {}
    try:
        for _, row in comentarios_df.iterrows():
            tipo = row['Tipo_comentario']
            comentario = row['Comentario']
            comentarios[tipo] = comentario
        return comentarios
    except Exception as e:
        st.error(f"Error al obtener comentarios: {e}")
        return {}
    
# Función para llenar la plantilla de Word
def llenar_plantilla_word(
    archivo_entrada, 
    archivo_salida, 
    valores, 
    grafico_path=None, 
    df_descripcion=None,
    df_total=None, 
    df_especies=None, 
    df_procedencia=None, 
    uploaded_images=None, 
    imagenes_dir=None
):
    try:
        # Cargar la plantilla de Word con docxtpl
        doc = DocxTemplate(archivo_entrada)
        
        # Si hay un gráfico para insertar, agrégalo al contexto
        if grafico_path:
            valores['grafico'] = InlineImage(doc, grafico_path, width=Mm(130))  # Ajusta el tamaño según sea necesario

        # Verificar los valores antes de renderizar
        logging.info("Valores para renderizar la plantilla:")
        for key, value in valores.items():
            if isinstance(value, InlineImage):
                logging.info(f"{key}: [InlineImage object]")
            elif isinstance(value, str):
                logging.info(f"{key}: {value[:50]}...")
            else:
                logging.info(f"{key}: {value}")

        #st.write("Contenido de 'valores' antes de renderizar la plantilla:", valores)
        
        # Preparar el diccionario para renderizar, preservando InlineImage
        valores_render = {}
        for k, v in valores.items():
            if isinstance(v, InlineImage):
                valores_render[k] = v
            elif isinstance(v, str):
                valores_render[k] = v
            else:
                valores_render[k] = str(v)
        
        # Renderizar la plantilla con los valores proporcionados
        doc.render(valores_render)
        
        # Guardar el documento generado
        doc.save(archivo_salida)
        st.success(f"Documento generado: {archivo_salida}")
        logging.info(f"Documento {archivo_salida} generado correctamente con los placeholders reemplazados.")

        # Abrir el documento con python-docx para llenar tablas e insertar imágenes
        doc_py = Document(archivo_salida)

        # Verificar que los DataFrames no sean None
        if df_descripcion is None:
            logging.error("Linea 1524: El DataFrame 'df_descripcion' es None.")
            return
        if df_total is None:
            logging.error("Linea 1527: El DataFrame 'df_total' es None.")
            return
        if df_especies is None:
            logging.error("Linea 1530: El DataFrame 'df_especies' es None.")
            return
        if df_procedencia is None:
            logging.error("Linea 1533: El DataFrame 'df_procedencia' es None.")
            return

        # Llenar las tablas y agregar imágenes
        llenar_primera_tabla(doc_py, df_descripcion)
        llenar_segunda_tabla(doc_py, df_total)
        llenar_tercera_tabla(doc_py, df_especies)
        llenar_tablas_procedencia(doc_py, df_procedencia)
        add_species_images_and_captions(doc_py, uploaded_images, imagenes_dir, df_especies)

        # Guardar los cambios en el documento
        doc_py.save(archivo_salida)
        st.success("Tablas e imágenes rellenadas correctamente.")
        logging.info("Tablas e imágenes rellenadas correctamente en el documento.")

    except Exception as e:
        st.error(f"Error al llenar la plantilla de Word: {e}")
        logging.error(f"Error al llenar la plantilla de Word: {e}", exc_info=True)

def leer_y_formatear_dataframes(excel_bytes_io):
    try:
        # Verificar que excel_bytes_io tiene datos
        excel_bytes_io.seek(0, os.SEEK_END)
        size = excel_bytes_io.tell()
        if size == 0:
            st.error("El archivo data.xlsx está vacío.")
            return ""
        excel_bytes_io.seek(0)  # Reiniciar el cursor
        
        logging.info(f"Tipo de excel_bytes_io: {type(excel_bytes_io)}")
        excel_file = pd.ExcelFile(excel_bytes_io, engine='openpyxl')
        logging.info("ExcelFile creado correctamente")

        # Leer cada hoja específica
        descripcion_df = excel_file.parse(sheet_name="descripcion")
        total_df = excel_file.parse(sheet_name="total")
        especies_df = excel_file.parse(sheet_name="especies")
        procedencia_df = excel_file.parse(sheet_name="procedencia")
        comentarios_df = excel_file.parse(sheet_name="comentarios")
        grafico_df = excel_file.parse(sheet_name="grafico")
        
        # Convertir cada DataFrame a texto
        descripcion_text = descripcion_df.to_string(index=False)
        total_text = total_df.to_string(index=False)
        especies_text = especies_df.to_string(index=False)
        procedencia_text = procedencia_df.to_string(index=False)
        comentarios_text = comentarios_df.to_string(index=False)
        grafico_text = grafico_df.to_string(index=False)
        
        # Combinar todos los textos en una sola cadena
        dataframes_text = f"""
        Descripción:
        {descripcion_text}

        Totales:
        {total_text}

        Especies:
        {especies_text}

        Procedencia:
        {procedencia_text}

        Comentarios:
        {comentarios_text}

        Gráfico:
        {grafico_text}
        """
        return dataframes_text
    except Exception as e:
        st.error(f"Error al leer y formatear data.xlsx: {e}")
        logging.error(f"Error al leer y formatear data.xlsx: {e}")
        return ""
        
def crear_donation_footer(base_dir):
    footer = st.container()
    
    with footer:
        st.markdown("---")
        st.header("💝 Apoya este proyecto")
        
        # Tabs for different payment methods
        tab1, tab2, tab3, tab4, tab5 = st.tabs(["Yape", "Depósito Bancario", "Tarjeta", "Otros Métodos", "Crypto"])
        
        # Yape Tab
        with tab1:
            col1, col2 = st.columns([1, 1])
            with col1:
                st.subheader("Donar por Yape")
                yape_image_path = os.path.join(base_dir, "yape.png")
                if os.path.exists(yape_image_path):
                    st.image(yape_image_path, width=300)
                else:
                    st.error(f"No se encontró la imagen en: {yape_image_path}")
        
        # Bank Deposits Tab
        with tab2:
            col1, col2 = st.columns([1, 1])
            
            with col1:
                st.subheader("En Soles 🇵🇪")
                banks_soles = {
                    "BCP": "31004315283063",
                    "BanBif": "008023869670",
                    "Interbank": "8983223094904",
                    "Scotiabank": "7508188435"
                }
                
                for bank, account in banks_soles.items():
                    st.write(f"**{bank}:**")
                    st_copy_to_clipboard(account, f"Copiar cuenta {bank}")
            
            with col2:
                st.subheader("En Dólares 💵")
                banks_usd = {
                    "BCP": "31004319160179",
                    "Interbank": "8983224537574",
                    "Scotiabank": "7508188145"
                }
                
                for bank, account in banks_usd.items():
                    st.write(f"**{bank}:**")
                    st_copy_to_clipboard(account, f"Copiar cuenta {bank}")
        
        # Card Payments Tab
        with tab3:
            st.subheader("Donar con tarjeta 💳")
            
            amounts = {
                "10": "https://pago-seguro.vendemas.com.pe/MTYzNDc3OTY3NjYxMWM4MjU2MTIuNzMxNzMxMjgxNTUz",
                "15": "https://pago-seguro.vendemas.com.pe/ZjE5MTc2MWIzMjM0MDQ3NDQ0NC4yOWQxNzMxMjgxNjIz",
                "20": "https://pago-seguro.vendemas.com.pe/NmEzOTkwNzI1OTY1Zi42MTM0MDg2NDYxNzMxMjgxNjUy",
                "25": "https://pago-seguro.vendemas.com.pe/MTUzODVjYTM4NDIzZDgxNjMwLjI0NzcxNzMxMjgxNjc3",
                "30": "https://pago-seguro.vendemas.com.pe/ODM0LjIyMTYyMjA2NjZmYjdhNmMzM2QxNzMxMjgxNzA3",
                "35": "https://pago-seguro.vendemas.com.pe/ODM3MzMzMDFkNDcuOTE2YzUzMjQxNTExNzMxMjgxNzI1",
                "40": "https://pago-seguro.vendemas.com.pe/YzgwZjM2NzM1LjZiMWQ0NzEzNTMxNzYxNzMxMjgxNzQ0",
                "45": "https://pago-seguro.vendemas.com.pe/MTgwMTYuNzQ1MzM0OTk0MzI4MTQ2MzYxNzMxMjgxNzYy"
            }
            
            cols = st.columns(4)
            for idx, (amount, link) in enumerate(amounts.items()):
                with cols[idx % 4]:
                    st.link_button(f"S/ {amount}", link)
            
            st.link_button("Más de S/ 50", "https://linkdecobro.ligo.live/v3/44df73097f594239b21b78b6905bed98")
        
        # Other Payment Methods Tab
        with tab4:
            col1, col2 = st.columns([1, 1])
            
            with col1:
                st.subheader("Mercado Pago")
                st.link_button("Donar con Mercado Pago", "https://link.mercadopago.com.pe/jersonapp")
            
            with col2:
                st.subheader("PayPal")
                st.link_button("Donar con PayPal", "https://www.paypal.com/paypalme/dschimbote")
        
        # Crypto Tab
        with tab5:
            st.subheader("Binance")
            col1, col2 = st.columns([1, 1])
            
            with col1:
                binance_image_path = os.path.join(base_dir, "binance.png")
                if os.path.exists(binance_image_path):
                    st.image(binance_image_path, width=300)
                else:
                    st.error(f"No se encontró la imagen en: {binance_image_path}")
            
            with col2:
                st.link_button("Donar con Binance", "https://app.binance.com/qr/dplkbb7f88c5329c4692adf278670d1b37ab")
                
def main():
    st.title("Genera tu Entregable")

    # Inicializar 'recomendaciones', 'data_excel', 'servicio' y actividades en session_state si no existen
    if 'recomendaciones' not in st.session_state:
        st.session_state['recomendaciones'] = ""
    if 'data_excel' not in st.session_state:
        st.session_state['data_excel'] = None
        logging.info(f"Tipo de st.session_state['data_excel']: {type(st.session_state['data_excel'])}")
    if 'servicio' not in st.session_state:
        st.session_state['servicio'] = ""
    if 'actividades' not in st.session_state:
        st.session_state['actividades'] = {"a": "", "b": "", "c": "", "d": "", "e": ""}

    # Subida de archivos Excel
    st.header("Sube tus reportes Excel")
    excel_files = st.file_uploader(
        "Selecciona tus archivos Excel",
        type=["xlsx", "xls"],
        accept_multiple_files=True,
    )

    # Subida de PDF
    st.header("Sube tu TDR (PDF)")
    pdf_file = st.file_uploader("Selecciona tu archivo PDF", type=["pdf"])

    # Subida de imágenes con selección de especie
    st.header("Adjuntar fotos de las especies")
    uploaded_images = st.file_uploader(
        "Seleccionar fotos", type=["png", "jpg", "jpeg"], accept_multiple_files=True
    )

    especies = cargar_especies()
    nombres_especies = list(especies.keys())

    if uploaded_images:
        for i, image in enumerate(uploaded_images):
            col1, col2 = st.columns([2, 3])
            with col1:
                st.image(image, caption=f"Imagen {i+1}", width=200)
            with col2:
                # Extraer el nombre base de la imagen sin extensión y en minúsculas
                base_name = os.path.splitext(image.name)[0].lower()

                # Filtrar especies que contienen el nombre base
                filtered_especies = [
                    name for name in nombres_especies if base_name in name.lower()
                ]
                # Ordenar alfabéticamente las especies filtradas
                sorted_filtered = sorted(filtered_especies)

                # Obtener el resto de las especies que no contienen el nombre base
                remaining_especies = [
                    name for name in nombres_especies if base_name not in name.lower()
                ]
                # Ordenar alfabéticamente las especies restantes
                sorted_remaining = sorted(remaining_especies)

                # Combinar las listas: primero las filtradas, luego las restantes
                combined_especies = sorted_filtered + sorted_remaining

                # Crear un key único para cada selectbox
                search_key = f"search_{i}"
                selected_especie = st.selectbox(
                    "Selecciona la especie:", combined_especies, key=search_key
                )
                st.write(f"Nombre científico: {especies[selected_especie]}")

    # Formulario para datos del entregable
    st.header("Datos del entregable")
    # Agregar nuevos campos al formulario
    dni = st.text_input("Introduce tu DNI", max_chars=8)
    # Variables para almacenar datos obtenidos de la API
    nombres = ""
    iniciales = ""
    ruc = ""
    # Botón para obtener datos de SUNAT
    if dni:
        if len(dni) == 8:
            nombres, iniciales, ruc = obtener_datos_sunat(dni)
            if nombres and iniciales and ruc:
                st.write(f"**Nombre Completo:** {nombres}")
                st.write(f"**Iniciales:** {iniciales}")
                st.write(f"**RUC:** {ruc}")
        else:
            st.error("El DNI debe tener exactamente 8 dígitos.")
    os_input = st.text_input("Introduce tu OS (4 dígitos)", max_chars=4)        
    email = st.text_input("Introduce tu Email", max_chars=50)
    telefono = st.text_input("Teléfono", max_chars=15)
    direccion = st.text_input("Dirección", max_chars=100)
    # Validaciones básicas
    if dni and not dni.isdigit():
        st.error("El DNI debe contener solo números.")
    if telefono and not telefono.isdigit():
        st.error("El teléfono debe contener solo números.")
    if os_input and not os_input.isdigit():
        st.error("La OS debe contener solo números.")
    if os_input and len(os_input) > 4:
        st.error("La OS debe tener como máximo 4 dígitos.")
    # Lista desplegable de tipos de entregable
    tipo_entregable = st.selectbox(
        "Selecciona el tipo de entregable",
        ["Unico", "Primer", "Segundo", "Tercer", "Cuarto"],
    )
    # Asignar valores a variables
    if tipo_entregable in ["Unico", "Primer"]:
        n = 1
    elif tipo_entregable == "Segundo":
        n = 2
    elif tipo_entregable == "Tercer":
        n = 3
    else:
        n = 4
    # Reemplazo de selección de mes y día por un calendario
    fecha_seleccionada = st.date_input(
        "Selecciona la fecha de presentación",
        value=datetime.now(),
        min_value=datetime(2000, 1, 1),
        max_value=datetime(2100, 12, 31)
    )

    # Extraer mes, día y año de la fecha seleccionada
    mes_num = fecha_seleccionada.month
    mes_seleccionado = MESES_ES.get(mes_num, fecha_seleccionada.strftime("%B"))
    dia_seleccionado = fecha_seleccionada.day
    year = fecha_seleccionada.year

    # Implementación en el formulario
    st.subheader("Selecciona la Ciudad y Departamento en el que desempeñas tus funciones")
    # Cargar el archivo ciudades.json
    ciudades = cargar_ciudades()
    departamentos = list(ciudades.keys()) + ["Ingreso manual"]
    departamento = st.selectbox("Selecciona el Departamento:", departamentos)
    # Mostrar ciudades o campo editable
    if departamento == "Ingreso manual":
        ciudad = st.text_input("Escribe el nombre de tu ciudad:")
    else:
        ciudad = st.selectbox("Selecciona la Ciudad:", ciudades[departamento])
    st.write(f"Departamento seleccionado: {departamento}, Ciudad seleccionada: {ciudad}")
    # Lógica para seleccionar el banco y autorrellenar el CCI en función del banco y cuenta
    st.subheader("Datos Bancarios")
    st.write("_Los datos bancarios son opcionales, solo se utilizan para generar el_ Formato 1 DJ Menores a 8 UIT")
    st.write("_De no ingresar sus datos bancarios estos no seran reemplazados en_ Formato 1 DJ Menores a 8 UIT")
    banco_seleccionado = st.selectbox("Selecciona tu banco:", ["BCP", "Interbank", "Scotiabank", "Banco de la Nación", "BanBif", "Otros"])
    cuenta = st.text_input("Ingresa tu cuenta", max_chars=20, placeholder="Ej. 123456789 o 123-456-78")

    # Generar el CCI basado en el banco y cuenta ingresada, si corresponde
    cci = ""
    if banco_seleccionado != "Otros" and cuenta:
        cuenta_limpia = cuenta.replace("-", "")
        if banco_seleccionado == "BCP":
            cci = "002" + cuenta_limpia + "13"
        elif banco_seleccionado == "Interbank":
            cci = "003" + cuenta_limpia + "43"
        elif banco_seleccionado == "Scotiabank":
            cci = "00936020" + cuenta_limpia + "95"
        elif banco_seleccionado == "Banco de la Nación":
            cci = "018-781-0" + cuenta_limpia + "-55"
        elif banco_seleccionado == "BanBif":
            cci = "0386501" + cuenta_limpia + "83"

    # Mostrar el CCI en un campo editable
    cci = st.text_input("CCI (editable):", value=cci)
    st.write(f"Banco seleccionado: {banco_seleccionado}, CCI generado: {cci}")

    # Subheader para las recomendaciones
    st.subheader("Ingrese su recomendación al Informe")

    # Lista desplegable para seleccionar el tipo de recomendación
    opcion_recomendacion = st.selectbox(
        "Selecciona el tipo de recomendación:",
        ["Generar recomendaciones por IA", "Ingresar recomendaciones manualmente"]
    )

    if opcion_recomendacion == "Generar recomendaciones por IA":
        # Botón para generar recomendaciones
        if st.button("Generar recomendaciones por IA"):
            # Validaciones previas
            if not excel_files:
                st.error("Por favor, sube al menos un archivo Excel antes de generar recomendaciones.")
            elif not pdf_file:
                st.error("Por favor, sube el archivo PDF de TDR antes de generar recomendaciones.")
            else:
                try:
                    # Crear directorios temporales para almacenar archivos
                    with tempfile.TemporaryDirectory() as carpeta_data:
                        # Guardar los archivos Excel subidos en la carpeta de datos
                        for file in excel_files:
                            with open(os.path.join(carpeta_data, file.name), "wb") as f:
                                f.write(file.getbuffer())
                        st.success("Archivos Excel guardados correctamente.")

                        # Procesar los archivos Excel para generar data.xlsx
                        output_excel = os.path.join(carpeta_data, "data.xlsx")
                        success = generar_data_excel(
                            [os.path.join(carpeta_data, file.name) for file in excel_files],
                            output_excel,
                        )
                        if success:
                            st.success("Archivo data.xlsx generado correctamente.")
                            # Leer data.xlsx y almacenarlo en session_state
                            with open(output_excel, "rb") as f:
                                st.session_state['data_excel'] = f.read()
                            st.success("Archivo data.xlsx almacenado en la sesión correctamente.")
                        else:
                            st.error("No se pudo generar data.xlsx.")

                        # Asegurarse de que data_excel no es None y es de tipo bytes
                        if st.session_state['data_excel'] and isinstance(st.session_state['data_excel'], bytes):
                            with open(output_excel, "rb") as f:
                                st.session_state['data_excel'] = f.read()
                                logging.info(f"Tipo de st.session_state['data_excel']: {type(st.session_state['data_excel'])}")
                        else:
                            st.error("El archivo data.xlsx no se generó correctamente o no está en el formato esperado.")

                        # Guardar data.xlsx en session_state como BytesIO
                        with open(output_excel, "rb") as f:
                            st.session_state['data_excel'] = f.read()
                            logging.info(f"Tipo de st.session_state['data_excel']: {type(st.session_state['data_excel'])}")

                        # **Guardar el PDF subido en un archivo temporal y definir tmp_pdf_path**
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                            tmp_pdf.write(pdf_file.getbuffer())
                            tmp_pdf_path = tmp_pdf.name
                        st.success("PDF de TDR guardado correctamente.")

                        # Abrir el archivo temporal en modo binario
                        with open(tmp_pdf_path, "rb") as f_pdf:
                            servicio_extraido = extraer_nombre_servicio(f_pdf)
                            f_pdf.seek(0)  # Reiniciar el cursor para la siguiente lectura
                            actividades_extraidas = extraer_actividades(f_pdf)
                            if servicio_extraido and actividades_extraidas:
                                st.session_state['servicio'] = servicio_extraido
                                st.session_state['actividades'] = {
                                    "a": actividades_extraidas[0],
                                    "b": actividades_extraidas[1],
                                    "c": actividades_extraidas[2],
                                    "d": actividades_extraidas[3],
                                    "e": actividades_extraidas[4],
                                }
                                logging.info("Datos de servicio y actividades extraídos correctamente.")
                            else:
                                st.error("Error al extraer datos de servicio y actividades del PDF.")

                        # **Eliminar el archivo temporal después de usarlo (opcional)**
                        #os.remove(tmp_pdf_path)

                        # Guardar 'servicio' y actividades en session_state
                        st.session_state['servicio'] = servicio_extraido
                        st.session_state['actividades'] = {
                            "a": actividades_extraidas[0],
                            "b": actividades_extraidas[1],
                            "c": actividades_extraidas[2],
                            "d": actividades_extraidas[3],
                            "e": actividades_extraidas[4],
                        }
                        # Después de extraer los datos
                        if servicio_extraido and all(isinstance(v, str) for v in actividades_extraidas):
                            st.session_state['servicio'] = servicio_extraido
                            st.session_state['actividades'] = {
                                "a": actividades_extraidas[0],
                                "b": actividades_extraidas[1],
                                "c": actividades_extraidas[2],
                                "d": actividades_extraidas[3],
                                "e": actividades_extraidas[4],
                            }
                            logging.info(f"Servicio guardado: {servicio_extraido}")
                            logging.info(f"Actividades guardadas: {st.session_state['actividades']}")

                        # Leer y formatear los dataframes
                        dataframes_text = leer_y_formatear_dataframes(BytesIO(st.session_state['data_excel']))

                        # Crear el prompt para la IA
                        prompt = f"""
                        Basándote en los siguientes datos, genera recomendaciones para el informe sin repetir la información proporcionada.

                        Datos Principales:
                        Servicio: {st.session_state['servicio']}
                        Actividades:
                        A: {st.session_state['actividades']['a']}
                        B: {st.session_state['actividades']['b']}
                        C: {st.session_state['actividades']['c']}
                        D: {st.session_state['actividades']['d']}
                        E: {st.session_state['actividades']['e']}

                        Datos Adicionales:
                        {dataframes_text}

                        Por favor, elabora recomendaciones basadas en esta información.
                        """

                        # Generar la respuesta de la IA
                        recomendaciones_ia = generar_respuesta_rapidapi(prompt)

                        # Almacenar las recomendaciones en session_state
                        st.session_state['recomendaciones'] = recomendaciones_ia

                        st.success("Recomendaciones generadas exitosamente.")

                except Exception as e:
                    st.error(f"Error al generar recomendaciones por IA: {e}")

        # Mostrar el campo de texto editable con las recomendaciones generadas por la IA
        recomendaciones = st.text_area(
            "Recomendaciones Generadas por IA (puedes editar):",
            value=st.session_state.get('recomendaciones', ''),
            height=200
        )
        # Actualizar session_state con cualquier edición del usuario
        st.session_state['recomendaciones'] = recomendaciones

    else:
        # Campo editable con la recomendación manual prellenada
        recomendaciones_manual = (
            "Coordinar con Capitanía de Puerto, respecto al caso de las matrículas y permisos "
            "de pesca de las embarcaciones pesqueras artesanales que operan en la zona, para que "
            "se formalicen pues hay un porcentaje de las mismas que aún no cuentan con estos documentos, "
            "pese a las amnistías que se les ha otorgado."
        )
        recomendaciones = st.text_area(
            "Ingrese sus recomendaciones manualmente:",
            value=recomendaciones_manual,
            height=200
        )
        # Almacenar las recomendaciones manuales en session_state
        st.session_state['recomendaciones'] = recomendaciones

    # Variables para almacenar datos extraídos del PDF
    servicio = ""
    act_a = act_b = act_c = act_d = act_e = ""

    # Botón para generar el entregable
    if st.button("Generar entregable"):
        # Validaciones
        if not excel_files:
            st.error("Por favor, sube al menos un archivo Excel para continuar.")
        elif not (telefono and dni and email and os_input and direccion):
            st.error("Por favor, completa todos los campos del formulario.")
        elif not ruc:
            st.error("No se pudo obtener el RUC. Verifica el DNI ingresado.")
        elif not pdf_file:
            st.error("Por favor, sube el archivo PDF de TDR.")
        elif not uploaded_images:
            st.error("Por favor, sube al menos una imagen.")
        elif not st.session_state['recomendaciones'].strip():
            st.error("Las recomendaciones no pueden estar vacías. Por favor, genera o ingresa recomendaciones.")
        else:
            try:
                # Crear directorios temporales para almacenar archivos
                with tempfile.TemporaryDirectory() as carpeta_data, tempfile.TemporaryDirectory() as carpeta_recursos:
                    # Guardar los archivos Excel subidos en la carpeta de datos
                    for file in excel_files:
                        with open(os.path.join(carpeta_data, file.name), "wb") as f:
                            f.write(file.getbuffer())
                    st.success("Archivos Excel guardados correctamente.")

                    # Verificar si data.xlsx ya fue generado al generar recomendaciones
                    if st.session_state['data_excel']:
                        logging.info(f"Tipo de st.session_state['data_excel']: {type(st.session_state['data_excel'])}")
                        # Guardar data.xlsx desde session_state
                        output_excel = os.path.join(carpeta_data, "data.xlsx")
                        with open(output_excel, "wb") as f:
                            f.write(st.session_state['data_excel'])
                        st.success("Archivo data.xlsx recuperado desde recomendaciones generadas.")
                    else:
                        # Generar data.xlsx si no existe
                        output_excel = os.path.join(carpeta_data, "data.xlsx")
                        success = generar_data_excel(
                            [os.path.join(carpeta_data, file.name) for file in excel_files],
                            output_excel,
                        )
                        if success:
                            st.success("Archivo data.xlsx generado correctamente.")
                            with open(output_excel, "rb") as f:
                                st.session_state['data_excel'] = f.read()
                            st.success("Archivo data.xlsx almacenado en la sesión correctamente.")
                        else:
                            st.error("No se pudo generar data.xlsx.")

                    # Copiar data.xlsx a carpeta_recursos
                    shutil.copy(output_excel, os.path.join(carpeta_recursos, "data.xlsx"))
                    st.success("Archivo data.xlsx copiado correctamente a la carpeta de recursos.")

                    # Guardar las imágenes subidas
                    for image in uploaded_images:
                        with open(os.path.join(carpeta_recursos, image.name), "wb") as f:
                            f.write(image.getbuffer())
                    st.success("Imágenes guardadas correctamente.")

                    # Guardar el PDF subido en un archivo temporal
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                        tmp_pdf.write(pdf_file.getbuffer())
                        pdf_path = tmp_pdf.name
                    st.success("PDF de TDR guardado correctamente.")

                    # Extraer datos del PDF
                    st.header("Extrayendo datos del PDF")
                    # Abrir el archivo temporal en modo binario
                    with open(pdf_path, "rb") as f_pdf:
                        servicio_extraido = extraer_nombre_servicio(f_pdf)
                        f_pdf.seek(0)  # Reiniciar el cursor para la siguiente lectura
                        actividades_extraidas = extraer_actividades(f_pdf)

                    # Contar el número de reportes Excel
                    reportes = contar_archivos(excel_files)
                    st.write(f"**Número de Reportes Excel:** {reportes}")

                    # Formatear la OS para que tenga 4 dígitos
                    orden = os_input.zfill(4)
                    st.write(f"**Orden de Servicio (OS):** {orden}")

                    # Generar el gráfico de barras
                    try:
                        df_grafico = pd.read_excel(output_excel, sheet_name="grafico")
                        grafico_path = generar_grafico(df_grafico, carpeta_recursos)
                        st.image(
                            grafico_path,
                            caption="Gráfico de Barras Generado",
                            use_column_width=True,
                        )
                        st.success("Gráfico generado correctamente.")
                    except Exception as e:
                        st.error(f"No se pudo generar el gráfico: {e}")
                        grafico_path = None

                    # Definir los valores para reemplazar en el documento
                    valores = {
                        "nombres": nombres,
                        "iniciales": iniciales,
                        "ruc": ruc,
                        "telefono": telefono,
                        "email": email,
                        "dni": dni,
                        "direccion": direccion,
                        "dia": dia_seleccionado,
                        "mes": mes_seleccionado,
                        "year": year,
                        "n": n,
                        "entregable": tipo_entregable,
                        "orden": orden,
                        "reportes": reportes,
                        "servicio": st.session_state.get('servicio', ''),
                        "act_a": st.session_state.get('actividades', {}).get('a', ''),
                        "act_b": st.session_state.get('actividades', {}).get('b', ''),
                        "act_c": st.session_state.get('actividades', {}).get('c', ''),
                        "act_d": st.session_state.get('actividades', {}).get('d', ''),
                        "act_e": st.session_state.get('actividades', {}).get('e', ''),
                        "ciudad": ciudad,
                        "banco": banco_seleccionado,
                        "cuenta": cuenta,
                        "cci": cci,
                        "recomendaciones": st.session_state.get('recomendaciones', ''),  # Incluir recomendaciones
                    }
                    
                    # Leer y formatear los dataframes
                    if st.session_state['data_excel']:
                        excel_bytes_io = BytesIO(st.session_state['data_excel'])
                        descripcion_df = pd.read_excel(excel_bytes_io, sheet_name="descripcion", engine='openpyxl')
                        total_df = pd.read_excel(excel_bytes_io, sheet_name="total", engine='openpyxl')
                        especies_df = pd.read_excel(excel_bytes_io, sheet_name="especies", engine='openpyxl')
                        procedencia_df = pd.read_excel(excel_bytes_io, sheet_name="procedencia", engine='openpyxl')
                        comentarios_df = pd.read_excel(excel_bytes_io, sheet_name="comentarios", engine='openpyxl')
                    else:
                        st.error("El archivo data.xlsx no está disponible en la sesión.")
                        return

                    # Verificar contenido de los DataFrames
                    logging.info("Contenido de 'descripcion_df':")
                    logging.info(descripcion_df.head())
                    st.write("DESARROLLO DE ACTIVIDADES:", descripcion_df[['Fecha', 'Hora', 'Descripcion']].head())
                    logging.info("Contenido de 'total_df':")
                    logging.info(total_df.head())
                    st.write("PRINCIPALES ESPECIES DESEMBARCADAS:", total_df.head())
                    logging.info("Contenido de 'especies_df':")
                    logging.info(especies_df.head())
                    st.write("VARIACION DE PRECIO DE LAS PRINCIPALES ESPECIES DESCARGAS:", especies_df[['NOMBRE COMUN', 'NOMBRE CIENTIFICO', 'PRECIO MAS BAJO (S/)', 'PRECIO MAS ALTO (S/)']].head())
                    logging.info("Contenido de 'procedencia_df':")
                    logging.info(procedencia_df.head())
                    st.write("Las zonas de pesca fueron:", procedencia_df.head())
                    # Obtener los comentarios de la hoja 'comentarios'
                    comentarios = obtener_comentarios_excel(comentarios_df)

                    # Actualizar los valores para reemplazar
                    valores.update({
                        'servicio': servicio_extraido,  # Usar el valor extraído
                        'act_a': actividades_extraidas[0],
                        'act_b': actividades_extraidas[1],
                        'act_c': actividades_extraidas[2],
                        'act_d': actividades_extraidas[3],
                        'act_e': actividades_extraidas[4],
                        'Especie_principal_1': comentarios.get('Especie_principal_1', ''),
                        'Especie_principal_2': comentarios.get('Especie_principal_2', ''),
                        'Especie_principal_3': comentarios.get('Especie_principal_3', ''), 
                        'Especie_minima_1': comentarios.get('Especie_minima_1', ''),
                        'Especie_minima_2': comentarios.get('Especie_minima_2', ''),
                        'Especie_minima_3': comentarios.get('Especie_minima_3', ''),         
                        'Mayor_dia': comentarios.get('Mayor_dia', ''),
                        'Menor_dia': comentarios.get('Menor_dia', ''),
                        'Aparejos': comentarios.get('Aparejos', ''),
                        'Mayor_Precio_1': comentarios.get('Mayor_Precio_1', ''),
                        'Mayor_Precio_2': comentarios.get('Mayor_Precio_2', ''),
                        'Mayor_Precio_3': comentarios.get('Mayor_Precio_3', ''),
                        'Menor_Precio_1': comentarios.get('Menor_Precio_1', ''),
                        'Menor_Precio_2': comentarios.get('Menor_Precio_2', ''),
                        'Menor_Precio_3': comentarios.get('Menor_Precio_3', ''),
                        'Total': comentarios.get('Total', '')  
                    })

                    # Definir el nombre del archivo de salida
                    archivo_salida = f"7. Informe N° 00{n}-{year}-{iniciales}.docx"
                    archivo_entrada = os.path.join(base_dir, "Informe.docx")  # Plantilla en la raíz
                    ruta_salida = os.path.join(carpeta_recursos, archivo_salida)
                    #st.write("Valores antes de llenar la plantilla:", valores)

                    # Verificar si la plantilla de Word existe
                    if not os.path.exists(archivo_entrada):
                        st.error(f"La plantilla {archivo_entrada} no se encontró.")
                    else:
                        # Llenar la plantilla de Word usando docxtpl
                        llenar_plantilla_word(
                            archivo_entrada, ruta_salida, valores, grafico_path
                        )

                        # Convertir el documento de Word a PDF
                        #archivo_pdf_salida = archivo_salida.replace(".docx", ".pdf")
                        #ruta_pdf_salida = os.path.join(
                        #    carpeta_recursos, archivo_pdf_salida
                        # )
                        #convertir_a_pdf(ruta_salida, ruta_pdf_salida)

                        # Llenar la plantilla y las tablas
                        llenar_plantilla_word(
                            archivo_entrada=archivo_entrada,
                            archivo_salida=ruta_salida,
                            valores=valores,  # Asegúrate de que 'valores' contiene todas las claves necesarias
                            grafico_path=grafico_path,
                            df_descripcion=descripcion_df,
                            df_total=total_df,
                            df_especies=especies_df,
                            df_procedencia=procedencia_df,
                            uploaded_images=uploaded_images,
                            imagenes_dir=carpeta_recursos
                        )

                        # Convertir el documento de Word a PDF
                        archivo_pdf_salida = archivo_salida.replace(".docx", ".pdf")
                        ruta_pdf_salida = os.path.join(carpeta_recursos, archivo_pdf_salida)
                        convertir_a_pdf(ruta_salida, ruta_pdf_salida)

                        # Procesar y convertir documentos adicionales a PDF
                        documentos_adicionales = [
                            "4. Formato 1 DJ Menores a 8 UIT.docx",
                            "Cargo.docx",
                            "Carta de Presentacion.docx",
                        ]
                        procesar_documentos_adicionales(
                            documentos_adicionales, carpeta_recursos, valores, base_dir
                        )

                        # Empaquetar todos los archivos generados en un ZIP
                        archivos_para_zip = [
                            archivo_salida, 
                            archivo_pdf_salida, 
                            "data.xlsx"  # Añade data.xlsx aquí
                        ] + [
                            docx.replace(".docx", ".pdf") for docx in documentos_adicionales
                        ]
                        zip_buffer = crear_zip_archivos(archivos_para_zip, carpeta_recursos)

                        # Proveer enlaces de descarga
                        st.header("Descarga de Archivos Generados")
                        st.download_button(
                            label="Descargar Todos los Archivos Generados (ZIP)",
                            data=zip_buffer,
                            file_name="archivos_generados.zip",
                            mime="application/zip",
                        )

                        st.success("Entregable generado exitosamente.")

            except Exception as e:
                st.error(f"Ocurrió un error durante el procesamiento: {e}")
                st.exception(e)

    crear_donation_footer(base_dir)
                
if __name__ == "__main__":
    main()
